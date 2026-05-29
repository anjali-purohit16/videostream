from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/project_presentation_qa.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D7DE"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Heading 1", 18, "B91C1C"),
        ("Heading 2", 13, "111827"),
        ("Heading 3", 11, "374151"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VideoStream Project Presentation Q&A Guide")
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(185, 28, 28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Module-wise workflow, viva questions, functions, and developer concepts")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(75, 85, 99)

    doc.add_paragraph()
    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = callout.cell(0, 0)
    set_cell_shading(cell, "FEF2F2")
    set_cell_margins(cell, 180, 220, 180, 220)
    p = cell.paragraphs[0]
    r = p.add_run("How to use this document: ")
    r.bold = True
    p.add_run(
        "Read the workflow first, then revise the module questions. "
        "In presentation, answer with the exact file/module name and the reason behind the design."
    )
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_qa(doc, questions):
    for q, a in questions:
        p = doc.add_paragraph()
        q_run = p.add_run("Q. " + q)
        q_run.bold = True
        q_run.font.color.rgb = RGBColor(17, 24, 39)
        p = doc.add_paragraph()
        a_run = p.add_run("A. ")
        a_run.bold = True
        p.add_run(a)


def add_module_table(doc):
    doc.add_heading("1. Project Module Map", level=1)
    doc.add_paragraph(
        "The project follows a custom MVC-style PHP structure. public/index.php is the front controller, "
        "controllers handle requests, models handle database operations, and views render user/admin screens."
    )

    rows = [
        ("Routing", "public/index.php", "Reads URL/action, maps it to controller and method, handles 404."),
        ("Authentication", "AuthController, AuthModel", "User/admin login, registration, OTP, logout, suspended/deleted checks."),
        ("User Home", "HomeController, home.php", "Builds user dashboard data and renders the user panel."),
        ("User Pages", "views/user/pages", "Browse, Trending, Categories, Watchlist, History, Profile, Subscription."),
        ("Video Playback", "PlaybackController, VideoModel, user.js", "Open modal, access check, record view, save progress."),
        ("Wishlist", "WishlistController, UserModel", "Add/remove video from user watchlist."),
        ("Review and Report", "UserFeedbackController, ReviewModel, ReportModel", "Submit review/rating and report video to admin."),
        ("Profile", "ProfileController, UserModel", "Update profile, password validation, delete account."),
        ("Subscription", "UserSubscriptionController, MessageModel, UserModel", "User sends plan request to admin."),
        ("Admin Dashboard", "DashboardController, Admin views", "Shows stats, charts, notifications, messages."),
        ("Admin Videos", "VideoController, VideoModel", "Upload, edit, delete, categorize videos."),
        ("Admin Users", "UserController, UserModel", "Create, suspend, activate, delete, export users."),
        ("Admin Review/Report", "ReviewController, ReportController", "Moderate reviews and handle reports."),
        ("Notifications", "NotificationModel, UserNotificationController, WsPublisher", "Display notifications and refresh via WebSocket."),
        ("Frontend Assets", "public/assets/css/user.css, public/assets/js/user.js", "Separated CSS/JS for user panel behavior and styling."),
        ("WebSocket", "websocket/server.php, app/core/WsPublisher.php", "Real-time event bridge for dashboard/user feed refresh."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["Module", "Main Files", "Responsibility"]
    widths = [1700, 2900, 4800]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_width(cell, widths[i])
        set_cell_shading(cell, "111827")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_width(cells[i], widths[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].paragraphs[0].add_run(text)
    set_table_borders(table)


def add_workflows(doc):
    doc.add_heading("2. Main Workflow of the Project", level=1)

    doc.add_heading("User Request Workflow", level=2)
    add_numbered(doc, [
        "Browser sends a URL such as /login, /profile, /movies, or ?action=save_review.",
        "public/index.php reads the URL and normalizes module, page, and action.",
        "The route table selects the correct controller, such as HomeController or ProfileController.",
        "Controller checks login/status using BaseController helpers and prepares data.",
        "Model runs database queries using PDO/prepared statements.",
        "Controller loads a view, and the layout includes CSS and user.js.",
        "For AJAX actions, controller returns JSON instead of a full page.",
    ])

    doc.add_heading("Login Workflow", level=2)
    add_numbered(doc, [
        "User enters email/password on login page.",
        "AuthController validates credentials using AuthModel/UserModel.",
        "If user is suspended or deleted, login is blocked.",
        "If valid, session data like user_id and user_name is stored in $_SESSION.",
        "Browser receives PHPSESSID cookie, which identifies the PHP session.",
        "User is redirected to the user panel.",
    ])

    doc.add_heading("User Panel Video Workflow", level=2)
    add_numbered(doc, [
        "HomeController loads videos, categories, history, wishlist, subscription, and notifications.",
        "home.php renders the shell and includes page partials like browse.php or profile.php.",
        "When a video card is clicked, user.js reads video data from the card payload.",
        "JS checks whether the user's plan can access the video.",
        "Video modal opens and playback begins.",
        "PlaybackController records views and saves progress through AJAX.",
    ])

    doc.add_heading("Admin Workflow", level=2)
    add_numbered(doc, [
        "Admin logs in through AuthController admin login.",
        "AdminController protects admin pages using requireAdmin.",
        "Admin views show dashboard, videos, categories, users, payments, reviews, reports, settings, and messages.",
        "Specific controllers handle each module, such as VideoController for videos and UserController for users.",
        "Actions update the database through models and create activity logs/notifications where needed.",
    ])

    doc.add_heading("Real-time Notification Workflow", level=2)
    add_numbered(doc, [
        "user.js opens a WebSocket connection to port 8080.",
        "Controllers call WsPublisher::push when something changes.",
        "WsPublisher sends an event to the bridge port 8081.",
        "websocket/server.php broadcasts the event to connected browsers.",
        "Browser calls feed_json to fetch latest data and rebuild notifications/feed.",
        "If WebSocket server is not running, data still updates after refresh or normal HTTP request.",
    ])


def add_module_qas(doc):
    doc.add_heading("3. Module-wise Questions and Answers", level=1)
    sections = {
        "Routing and MVC": [
            ("What is the role of public/index.php?", "It is the front controller. Every request enters here, then it maps the URL/action to the correct controller and method."),
            ("Why did we add route maps?", "Route maps make friendly URLs like /profile and /admin/videos easier to manage and also allow invalid URLs to show 404 instead of silently loading home."),
            ("What is MVC in this project?", "Model stores database logic, View renders HTML, Controller handles request flow between the user and the model/view."),
            ("Why split actions from HomeController?", "To keep each responsibility small. Wishlist, profile, playback, feedback, subscription, history, and notifications now have their own focused controllers."),
        ],
        "Authentication and Sessions": [
            ("Where is session data stored?", "The browser stores the PHPSESSID cookie. The actual session values like $_SESSION['user_id'] are stored server-side by PHP."),
            ("How do we stop suspended users from logging in?", "During authentication, the system checks user status. Suspended or deleted users are rejected and active sessions can be purged from admin actions."),
            ("Why is requireActiveUser important?", "It protects user-only actions from guests and inactive users before any sensitive operation runs."),
            ("What happens on logout?", "The session is cleared/destroyed, so the current PHPSESSID no longer gives access to logged-in pages."),
        ],
        "User Panel": [
            ("What does HomeController do now?", "It renders the home page, prepares shared user dashboard data, formats feed data, handles visible notifications, and returns feed_json."),
            ("Why is home.php still important?", "It is the main user-panel view shell. It includes sidebar, topbar, page partials, footer, video modal, and frontend config."),
            ("Why did we separate CSS and JS?", "It reduced confusion in home.php and made styling/behavior easier to maintain in user.css and user.js."),
            ("How are pages like profile/history loaded?", "Pretty URLs set upage, then home.php includes the matching partial from app/views/user/pages."),
        ],
        "Videos and Playback": [
            ("How does video access control work?", "Each video has an access level. JS and server-side data compare user's plan rank with video plan rank before opening premium/basic content."),
            ("How is watch progress saved?", "user.js sends progress through AJAX to PlaybackController::save_progress, which stores it using UserModel."),
            ("How is a view recorded?", "When a video opens/plays, user.js calls record_view and VideoModel increments view count."),
            ("Why is the video modal in an include file?", "Because it is shared by multiple user pages and keeping it separate makes home.php cleaner."),
        ],
        "Wishlist and History": [
            ("How does wishlist work?", "user.js sends wishlist_toggle or remove_wishlist. WishlistController calls UserModel to update the database and returns/redirects."),
            ("Why did we move wishlist DB queries into UserModel?", "Because database logic belongs in the model, not directly in the controller."),
            ("How does history clear work?", "HistoryController calls UserModel::clearWatchHistory for the logged-in user and redirects to history."),
            ("Why keep user_id in every history/wishlist query?", "To ensure one user cannot access or modify another user's records."),
        ],
        "Reviews and Reports": [
            ("What happens when user submits a review?", "UserFeedbackController validates video/rating/comment and ReviewModel saves or updates the review."),
            ("Why are reviews not instantly public?", "They can be stored with a status like pending so admin can approve or moderate them."),
            ("What happens when user reports a video?", "UserFeedbackController creates a report through ReportModel and notifies admin."),
            ("Why separate ReviewModel and ReportModel?", "They represent different database concepts and admin workflows."),
        ],
        "Profile and Subscription": [
            ("How does profile update work?", "ProfileController validates name/password fields and UserModel updates the profile or password hash."),
            ("Why validate current password before changing password?", "It confirms the real account owner is requesting a sensitive change."),
            ("How does delete account work?", "ProfileController handles the user request and updates/removes the account as per project logic."),
            ("How does subscription request work?", "UserSubscriptionController validates selected plan and creates a message/request for admin approval."),
        ],
        "Admin Panel": [
            ("What does AdminController provide?", "Common admin page setup: require admin, notifications, messages, flash data, and shared layout data."),
            ("How does admin video edit work?", "VideoController loads existing video data, the admin view shows edit form, and save updates VideoModel."),
            ("How does admin user suspend work?", "UserController updates status to suspended, purges active sessions, pushes real-time account status event, and redirects back."),
            ("Why have separate admin controllers?", "Admin modules are operationally different. Separate controllers make dashboard, videos, users, reports, reviews, and settings easier to maintain."),
        ],
        "Notifications and WebSocket": [
            ("Are user notifications real-time?", "Yes, when websocket/server.php is running. Otherwise they update through normal page refresh/feed request."),
            ("What is WsPublisher?", "A small bridge publisher that lets PHP controllers send events to the WebSocket server."),
            ("What is feed_json?", "A JSON endpoint in HomeController that returns fresh user panel data after a WebSocket event."),
            ("Why not send all data directly through WebSocket?", "The project uses WebSocket only as a signal. The browser fetches real data over HTTP, which keeps logic simpler and safer."),
        ],
    }
    for heading, qas in sections.items():
        doc.add_heading(heading, level=2)
        add_qa(doc, qas)


def add_presentation_questions(doc):
    doc.add_heading("4. Important Presentation/Viva Questions", level=1)
    qas = [
        ("Why did you choose MVC?", "MVC separates database work, request handling, and UI rendering, so the project is easier to debug and extend."),
        ("What problem did your refactoring solve?", "The user panel had too much CSS, JS, and logic in home.php. Refactoring separated views, scripts, styles, controllers, and model methods while preserving behavior."),
        ("How did you make sure functionality was not broken?", "We followed one small move, test. After each move, PHP syntax checks, JS checks, route checks, and endpoint checks were run."),
        ("How do you handle invalid URLs?", "public/index.php now routes known URLs and returns a 404 view for unknown paths instead of loading home for every random URL."),
        ("What is the difference between user panel and admin panel?", "User panel focuses on browsing/watching videos, wishlist, reviews, reports, profile, and subscription. Admin panel manages content, users, reviews, reports, payments, settings, and dashboard monitoring."),
        ("How do you secure user actions?", "Actions call requireActiveUser, use the session user_id, validate inputs, and perform database operations through model methods."),
        ("How is SQL injection reduced?", "Models use PDO prepared statements and bound parameters instead of directly concatenating user input into SQL."),
        ("How does real-time update work?", "Controllers publish an event through WsPublisher, WebSocket server broadcasts it, and the browser refreshes JSON data through feed_json."),
        ("What is the purpose of user.js?", "It contains user-panel frontend behavior: notifications, feed refresh, wishlist, subscription request, review/report, profile update, and video modal playback."),
        ("What is the purpose of user.css?", "It centralizes user-panel styling that was previously spread across view/style partials."),
        ("What was the riskiest part of refactoring?", "Video player/modal JavaScript, because it controls access checks, playback, progress, fullscreen, seek, volume, reviews, and reports."),
        ("What would you improve next?", "Add automated feature tests for logged-in flows, move remaining view helper functions out of home.php, and add clearer service classes for home data preparation."),
    ]
    add_qa(doc, qas)


def add_functions_section(doc):
    doc.add_heading("5. User-defined and Predefined Functions in This Project", level=1)
    doc.add_paragraph(
        "A user-defined function is written by the developer inside the project. A predefined function is already provided by PHP or JavaScript."
    )

    doc.add_heading("User-defined Functions/Methods", level=2)
    rows = [
        ("u_page_url()", "app/views/user/home.php", "Builds user-panel friendly page URLs."),
        ("u_video_payload()", "app/views/user/home.php", "Creates safe JSON payload used by video cards."),
        ("buildHomeViewData()", "HomeController.php", "Collects data needed to render the user home panel."),
        ("feed_json()", "HomeController.php", "Returns fresh user feed data as JSON."),
        ("toggleWishlist()", "UserModel.php", "Adds or removes a video from user wishlist."),
        ("saveWatchProgress()", "UserModel.php", "Stores video progress for the logged-in user."),
        ("saveUserReview()", "ReviewModel.php", "Creates or updates a user's review."),
        ("push()", "WsPublisher.php", "Sends a real-time topic event to the WebSocket bridge."),
        ("purgeUserSessions()", "UserController.php", "Logs out a suspended/deleted user by clearing sessions."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [2300, 2700, 4300]
    for i, text in enumerate(["Function/Method", "File", "Purpose"]):
        cell = table.rows[0].cells[i]
        set_width(cell, widths[i])
        set_cell_shading(cell, "111827")
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_width(cells[i], widths[i])
            set_cell_margins(cells[i])
            cells[i].paragraphs[0].add_run(text)
    set_table_borders(table)

    doc.add_heading("Predefined PHP/JS Functions Used", level=2)
    add_bullets(doc, [
        "session_start(), $_SESSION: starts and reads server-side PHP sessions.",
        "json_encode(), json_decode(): converts PHP arrays to JSON and JSON back to arrays.",
        "strtolower(), trim(), substr(), explode(): string handling.",
        "date(), strtotime(), time(): date/time formatting and timestamps.",
        "password_hash(), password_verify(): secure password storage and verification.",
        "filter_var(): input validation, commonly used for email validation.",
        "header(), http_response_code(): sends HTTP response headers/status.",
        "move_uploaded_file(): saves uploaded files.",
        "file_get_contents(): reads a file or remote response, used carefully.",
        "fetch(): JavaScript API for AJAX requests.",
        "addEventListener(): JavaScript event binding.",
        "setTimeout(), setInterval(): JavaScript timers for toast/progress logic.",
        "JSON.parse(), JSON.stringify(): JavaScript JSON handling.",
    ])


def add_concepts(doc):
    doc.add_heading("6. Important Developer Concepts to Know", level=1)
    concepts = [
        ("MVC Architecture", "Keeps the application maintainable by separating controller flow, model database logic, and view rendering."),
        ("Front Controller Pattern", "public/index.php acts as the single entry point and decides which controller should handle the request."),
        ("Sessions and Cookies", "PHPSESSID is stored in the browser cookie; actual user data is stored server-side in $_SESSION."),
        ("Authentication vs Authorization", "Authentication checks who the user is. Authorization checks what that user is allowed to access."),
        ("Prepared Statements", "Protect database queries by separating SQL structure from user-provided values."),
        ("CSRF Awareness", "Important for forms because a logged-in browser can be tricked into sending unwanted POST requests if CSRF protection is missing."),
        ("XSS Prevention", "Escape output using helper functions like h() before showing user/database content in HTML."),
        ("AJAX/JSON Endpoints", "Used for actions like review submit, report submit, wishlist, profile update, and feed refresh without full page reload."),
        ("WebSocket Signaling", "Useful for real-time updates. In this project, WebSocket tells browser to fetch latest data over HTTP."),
        ("Single Responsibility Principle", "Each controller/model method should do one focused job. This is why large HomeController logic was split."),
        ("Progressive Refactoring", "Move one small part, test, then continue. This reduces risk in an existing working project."),
        ("Error Handling", "Invalid routes should show 404; server errors should not expose sensitive details to users."),
        ("Asset Separation", "CSS belongs in stylesheets, JS belongs in script files, and PHP views should mainly render HTML/data."),
        ("Input Validation", "Never trust request data. Validate IDs, email, password, plan IDs, review rating, and report reason."),
        ("Role-based Access", "Admin and user panels must have different guards so normal users cannot access admin actions."),
    ]
    for title, body in concepts:
        p = doc.add_paragraph()
        r = p.add_run(title + ": ")
        r.bold = True
        r.font.color.rgb = RGBColor(185, 28, 28)
        p.add_run(body)


def add_pending_and_confidence(doc):
    doc.add_heading("7. Current Status and Suggested Next Improvements", level=1)
    add_bullets(doc, [
        "Routing, controller split, CSS extraction, JS extraction, and model query movement are mostly complete.",
        "HomeController is much cleaner and now focuses on home/feed data.",
        "home.php is organized with readable sections but still contains view helper functions and the PHP-to-JS config block.",
        "Manual logged-in testing should still cover video modal, playback, wishlist, review/report, profile, subscription, history clear, and notifications.",
        "Future improvement: move view helper functions from home.php into a dedicated helper file if the project grows further.",
        "Future improvement: add automated tests for main user/admin workflows.",
    ])


def main():
    doc = Document()
    style_doc(doc)
    add_title(doc)
    add_module_table(doc)
    add_workflows(doc)
    add_module_qas(doc)
    add_presentation_questions(doc)
    add_functions_section(doc)
    add_concepts(doc)
    add_pending_and_confidence(doc)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("VideoStream Project Q&A Guide")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(107, 114, 128)

    doc.save(OUT)


if __name__ == "__main__":
    main()
